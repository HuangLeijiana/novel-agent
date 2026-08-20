"""Serialization utilities for YAML, Markdown, and DOCX formats."""

from enum import Enum
from pathlib import Path
from typing import Any

import yaml
from pydantic import BaseModel

# ============================================================
# YAML
# ============================================================


class YamlSerializer:
    """Handles YAML serialization with sensible defaults."""

    @staticmethod
    def to_yaml(obj: Any, file_path: Path) -> None:
        """Serialize an object to a YAML file.

        Pydantic models are converted to dict first.
        """
        file_path.parent.mkdir(parents=True, exist_ok=True)
        data = YamlSerializer._to_dict(obj)
        with open(file_path, "w", encoding="utf-8") as f:
            yaml.dump(
                data,
                f,
                allow_unicode=True,
                default_flow_style=False,
                sort_keys=False,
                indent=2,
                width=120,
            )

    @staticmethod
    def from_yaml(file_path: Path, model_class: type[BaseModel] | None = None) -> Any:
        """Deserialize a YAML file to a dict or Pydantic model."""
        with open(file_path, encoding="utf-8") as f:
            data = yaml.safe_load(f)
        if model_class is not None and data is not None:
            return model_class.model_validate(data)
        return data

    @staticmethod
    def _to_dict(obj: Any) -> Any:
        """Convert an object to a dict suitable for YAML serialization."""
        if isinstance(obj, Enum):
            return obj.value
        if isinstance(obj, BaseModel):
            return YamlSerializer._to_dict(obj.model_dump(mode="python", exclude_none=False))
        if isinstance(obj, dict):
            return {k: YamlSerializer._to_dict(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [YamlSerializer._to_dict(item) for item in obj]
        if isinstance(obj, tuple):
            return list(obj)
        if isinstance(obj, Path):
            return str(obj)
        return obj


# ============================================================
# Markdown
# ============================================================


class MarkdownSerializer:
    """Handles Markdown file generation for chapters."""

    @staticmethod
    def chapter_to_markdown(
        chapter_number: int,
        title: str,
        content: str,
        include_frontmatter: bool = True,
    ) -> str:
        """Build a complete Markdown document for a chapter."""
        lines = []
        if include_frontmatter:
            lines.append("---")
            lines.append(f"chapter: {chapter_number}")
            if title:
                lines.append(f'title: "{title}"')
            lines.append("---")
            lines.append("")
        lines.append(f"# 第{chapter_number}章 {title}" if title else f"# 第{chapter_number}章")
        lines.append("")
        lines.append(content)
        return "\n".join(lines)

    @staticmethod
    def save_chapter_markdown(
        file_path: Path,
        chapter_number: int,
        title: str,
        content: str,
    ) -> None:
        """Save a chapter as a markdown file."""
        file_path.parent.mkdir(parents=True, exist_ok=True)
        md = MarkdownSerializer.chapter_to_markdown(chapter_number, title, content)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(md)

    @staticmethod
    def load_chapter_markdown(file_path: Path) -> tuple[int, str, str]:
        """Load a chapter from a markdown file. Returns (chapter_number, title, content)."""
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
        # Simple parsing: extract title from H1, return the rest as content
        lines = text.split("\n")
        title = ""
        content_start = 0
        for i, line in enumerate(lines):
            if line.startswith("# "):
                title = line.lstrip("# ").strip()
                content_start = i + 1
                break
        content = "\n".join(lines[content_start:]).strip()
        return 0, title, content


# ============================================================
# DOCX (Word Document)
# ============================================================


class DocxSerializer:
    """Handles Word document generation for chapters using python-docx."""

    @staticmethod
    def _normalize_text(text: str) -> str:
        """Replace Unicode compatibility symbols with their ASCII forms.

        Some fonts/viewers render U+2103 (℃) / U+2109 (℉) as blank glyphs,
        making the character invisible in Word/WPS even though the data is
        intact. °C / °F are composed from widely-supported characters.
        """
        return text.replace("℃", "°C").replace("℉", "°F")

    @staticmethod
    def _set_east_asian_font(doc) -> None:
        """Set 宋体 as the explicit East Asian font on used styles.

        python-docx's ``font.name`` only sets w:ascii/w:hAnsi (the Latin
        font). Without an explicit w:eastAsia, CJK text is left to the
        viewer's theme/default fallback font, which can render some
        characters as blanks.
        """
        from docx.oxml.ns import qn

        for style_name in ("Normal", "Title", "Heading 1"):
            try:
                st = doc.styles[style_name]
            except KeyError:
                continue
            rpr = st.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def chapter_to_docx(
        file_path: Path,
        chapter_number: int,
        title: str,
        content: str,
    ) -> None:
        """Save a chapter as a formatted .docx file."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        file_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        DocxSerializer._set_east_asian_font(doc)

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)

        # Title
        heading_text = f"第{chapter_number}章"
        if title:
            heading_text += f" {DocxSerializer._normalize_text(title)}"
        heading = doc.add_heading(heading_text, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a blank line after title
        doc.add_paragraph("")

        # Content — split on double newlines for paragraph breaks
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            # Handle inline line breaks
            sub_paras = para_text.split("\n")
            for j, sub in enumerate(sub_paras):
                sub = sub.strip()
                if not sub:
                    continue
                if sub.startswith("### "):
                    p = doc.add_paragraph()
                    run = p.add_run(DocxSerializer._normalize_text(sub[4:]))
                    run.bold = True
                    continue
                p = doc.add_paragraph(DocxSerializer._normalize_text(sub))
                # First line indent for Chinese text
                p.paragraph_format.first_line_indent = Cm(0.74)
                style = p.style
                style.font.size = Pt(12)
                style.font.name = "宋体"

        doc.save(str(file_path))

    @staticmethod
    def book_to_docx(
        file_path: Path,
        title: str,
        chapters: list[tuple[int, str, str]],
    ) -> None:
        """Export an entire book as a single .docx file.

        Args:
            file_path: Output path
            title: Book title
            chapters: List of (chapter_number, chapter_title, content) tuples
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt

        doc = Document()
        DocxSerializer._set_east_asian_font(doc)

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)

        # Book title page
        title_heading = doc.add_heading(DocxSerializer._normalize_text(title), level=0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # Chapters
        for chapter_number, chapter_title, content in chapters:
            heading_text = f"第{chapter_number}章"
            if chapter_title:
                heading_text += f" {DocxSerializer._normalize_text(chapter_title)}"
            doc.add_heading(heading_text, level=1)
            doc.add_paragraph("")

            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                sub_paras = para_text.split("\n")
                for sub in sub_paras:
                    sub = sub.strip()
                    if not sub:
                        continue
                    if sub.startswith("### "):
                        p = doc.add_paragraph()
                        run = p.add_run(DocxSerializer._normalize_text(sub[4:]))
                        run.bold = True
                        continue
                    p = doc.add_paragraph(DocxSerializer._normalize_text(sub))
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    style = p.style
                    style.font.size = Pt(12)
                    style.font.name = "宋体"

            # Page break after each chapter (except the last)
            if chapter_number != chapters[-1][0]:
                doc.add_page_break()

        doc.save(str(file_path))
